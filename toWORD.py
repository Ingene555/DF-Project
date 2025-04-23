""" 
This library will be used for conversion into word
A class per file type will be defined
"""

import win32com.client as client32
import os
from datetime import datetime as dt
import mainFunctions as mf
from docx import Document
from PyPDF2 import PdfReader
from PIL import Image
from io import BytesIO
from pyxlsb import open_workbook
from openpyxl import load_workbook
import csv
from toIMG import ModifyToImg
import subprocess as spss
import time

TEMPCACHES = 'tempcaches'


def pls():
    powershell_command = """
    $PowerPoint = New-Object -ComObject PowerPoint.Application
    $PowerPoint.Visible = -1  # msoTrue
    Write-Host "PowerPoint lancé avec succès."
    """
    try:
        result = spss.run(
            ["powershell", "-Command", powershell_command],
            capture_output=True, text=True, shell=True
        )
        return [True, result]
    except Exception as e:
        return [False, e]


def ptp(input_file, output_file):
    r = pls()
    if r[0]:
        powerpoint = client32.Dispatch("PowerPoint.Application")
        presentation = powerpoint.Presentations.Open(os.path.abspath(input_file))
        presentation.SaveAs(os.path.abspath(output_file), 24)
        presentation.Close()
        time.sleep(.1)
        return [True, output_file]
    else:return r


class WordToWord:

    def __init__(self, file):
        self.file = file if type(file) == list else [file]
        self.stop = False
        self.progressShow = None

    def toWORD(self, path):
        path = path if type(path) == list else [path]
        __ = list()
        for xx, ww in enumerate(self.file):
            if not self.stop:
                try:
                    word = client32.Dispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(os.path.abspath(ww))
                    doc.SaveAs(os.path.abspath(path[xx]), FileFormat=16)
                    doc.Close()
                    word.Quit()
                    __.append(mf.setReturn(True, [type(path[xx])], [path[xx]]))
                    time.sleep(.1)
                except Exception as e:
                    __.append(mf.setReturn(False, [type(e)], [e]))
                if self.progressShow:
                    self.progressShow(int(xx * 100 / len(self.file)))
            else:
                for i in __:
                    try:
                        os.remove(mf.Return(i).getValues(0))
                    except:pass
                __=[]
                break
        if self.progressShow:
            self.progressShow(100, True, self.stop, _)
        return __

    def kill(self):
        self.stop = True


class TextToWord:

    def __init__(self, file):
        self.file = file if type(file) == list else [file]
        self.stop = False
        self.toKill = None
        self.progressShow = None

    def Convert(self, pathF, pathT):
        pathT
        _ = list()
        for ww in range(len(pathF)):
            type(ww)
            _.append(os.path.join(TEMPCACHES, f"tc-textrd-{ww}-{dt.now()}.docx".replace(":", "")))
        while(len(pathF) > len(_)):
            _.append(os.path.join(TEMPCACHES, f"tc-textrd-{dt.now()}.docx".replace(":", "")))
        __ = list()
        for xx, ww in enumerate(pathF):
            if not self.stop:
                try:
                    with open(ww, 'r', encoding='utf-8') as f:
                        lines = f.readlines()
                    doc = Document()
                    for yy, line in enumerate(lines):
                        if not self.stop:
                            doc.add_paragraph(line)
                            if self.progressShow:
                                self.progressShow(int(yy * 100 / len(lines)))
                        else:
                            break
                    doc.save(_[xx])
                    __.append(mf.setReturn(True, [type(_[xx])], [_[xx]]))
                except Exception as e:
                    __.append(mf.setReturn(False, [type(e)], [e]))
            else:
                for i in __:
                    try:
                        os.remove(mf.Return(i).getValues(0))
                    except:pass
                break
        if self.progressShow:
            self.progressShow(100)
        return __

    def toWORD(self, path):
        temp = self.Convert(self.file, [])
        path = path if type(path) == list else path
        _=list()
        if not self.stop:
            self.toKill = WordToWord([mf.Return(ww).getValues(0) for ww in temp])
            _ = self.toKill.toWORD(path)
        if self.progressShow:
            self.progressShow(100, True, self.stop, _)
        return _ if not self.stop else []

    def kill(self):
        self.stop = True
        if self.toKill:
            try:
                self.toKill.kill()
            except Exception as e:
                print(e)


class MixedToWord:

    def __init__(self, *file):
        self.file = [*file]
        n = False
        for ww in self.file:
            if type(ww) != list:
                n = True
                break
        if n:self.file = [self.file]
        self.stop = False
        self.toKill = None
        self.progressShow = None

    def Convert(self, pathF, pathT):
        pathT
        _ = list()
        for ww in range(len(pathF)):
            type(ww)
            _.append(os.path.join(TEMPCACHES, f"tc-mdrd-{ww}-{dt.now()}.docx".replace(":", "")))
        while(len(pathF) > len(_)):
            _.append(os.path.join(TEMPCACHES, f"tc-mdrd-{dt.now()}.docx".replace(":", "")))
        __ = list()
        for xx, ww in enumerate(pathF):
            if not self.stop:
                try:
                    doc = Document()
                    for yy, item in enumerate(ww):
                        try:
                            if len(item) < 3: item.append(None)
                            item_type, content, none = item
                            if item_type == "text":
                                with open (content, "r", encoding="utf-8") as f:
                                    for zz in f.readlines():
                                        doc.add_paragraph(zz)
                            elif item_type == "string":
                                doc.add_paragraph(content)
                            elif item_type == "pdf":
                                with open(content, 'rb') as f:
                                    reader = PdfReader(f)
                                    pdf_text = ""
                                    for page in reader.pages:
                                        pdf_text += page.extract_text()
                                    doc.add_paragraph(pdf_text)
                                    doc.add_paragraph()
                            elif item_type == "image":
                                tm = f"{TEMPCACHES}/tc-{dt.now()}.png".replace(":", "")
                                self.toKill = ModifyToImg(content)
                                im = mf.Return(self.toKill.toIMG(tm, none if none else {})[0]).getValues(0)
                                img = Image.open(im)
                                img_byte_arr = BytesIO()
                                img.save(img_byte_arr, format="PNG")
                                img_byte_arr.seek(0)
                                doc.add_picture(img_byte_arr)
                                doc.add_paragraph()
                                try:os.remove(im)
                                except Exception as e:print(e)
                            elif item_type == "word":
                                tm = f"{TEMPCACHES}/tc-{dt.now()}.docx".replace(":", "")
                                self.toKill = WordToWord(content)
                                wd = mf.Return(self.toKill.toWORD(tm)[0]).getValues(0)
                                doc_to_read = Document(wd)
                                for para in doc_to_read.paragraphs:
                                    doc.add_paragraph(para.text)
                                    doc.add_paragraph()
                                try:os.remove(wd)
                                except Exception as e:print(e)
                            elif item_type == "excel":
                                tm = f"{TEMPCACHES}/tc-{dt.now()}.docx".replace(":", "")
                                self.toKill = ExcelToWord(content)
                                ex = mf.Return(self.toKill.toWORD(tm)[0]).getValues(0)
                                doc_to_read = Document(ex)
                                for para in doc_to_read.paragraphs:
                                    doc.add_paragraph(para.text)
                                    doc.add_paragraph()
                                try:os.remove(ex)
                                except Exception as e:print(e)
                        except Exception as e: print(e)
                        if self.progressShow:
                            self.progressShow(int(yy * 100 / len(ww)))
                    doc.save(_[xx])
                    __.append(mf.setReturn(True, [type(_[xx])], [_[xx]]))
                except Exception as e:
                    __.append(mf.setReturn(False, [type(e)], [e]))
            else:
                for i in __:
                    try:
                        os.remove(mf.Return(i).getValues(0))
                    except:pass
                break
        if self.progressShow:
            self.progressShow(100)
        return __

    def toWORD(self, path):
        temp = self.Convert(self.file, [])
        path = path if type(path) == list else path
        _ = list()
        if not self.stop:
            self.toKill = WordToWord([mf.Return(ww).getValues(0) for ww in temp])
            _ = self.toKill.toWORD(path)
        if self.progressShow:
            self.progressShow(100, True, self.stop, _)
        return _ if not self.stop else []

    def kill(self):
        self.stop = True
        if self.toKill:
            try:
                self.toKill.kill()
            except Exception as e:
                print(e)


class ExcelToWord:

    def __init__(self, file):
        self.file = file if type(file) == list else [file]
        self.stop = False
        self.toKill = None
        self.progressShow = None

    def Convert(self, pathF, pathT):
        pathT
        _ = list()
        for ww in range(len(pathF)):
            type(ww)
            _.append(os.path.join(TEMPCACHES, f"tc-excelrd-{ww}-{dt.now()}.docx".replace(":", "")))
        while(len(pathF) > len(_)):
            _.append(os.path.join(TEMPCACHES, f"tc-excelrd-{dt.now()}.docx".replace(":", "")))
        __ = list()
        for xx, ww in enumerate(pathF):
            if not self.stop:
                try:
                    doc = Document()
                    ext = os.path.splitext(ww)[-1].lower()
                    if ext == ".xlsb":
                        with open_workbook(ww) as wb:
                            for yy, sheet_name in enumerate(wb.sheets):
                                doc.add_heading(f"{sheet_name}", level=2)
                                with wb.get_sheet(sheet_name) as sheet:
                                    table_data = []
                                    for row in sheet.rows():
                                        table_data.append([item.v for item in row])
                                    if table_data:
                                        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                                        for i, row in enumerate(table_data):
                                            for j, value in enumerate(row):
                                                table.cell(i, j).text = str(value) if value is not None else ""
                                if self.progressShow:
                                    self.progressShow(int(yy * 100 / len(wb.sheets)))
                    elif ext in [".xls", ".xlsx", ".xlsm", ".xltx", ".xlt"]:
                        workbook = load_workbook(ww, data_only=True)
                        for yy, sheet_name in enumerate(workbook.sheetnames):
                            sheet = workbook[sheet_name]
                            doc.add_heading(f"{sheet_name}", level=2)
                            table_data = []
                            for row in sheet.iter_rows(values_only=True):
                                table_data.append(row)
                            if table_data:
                                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                                for i, row in enumerate(table_data):
                                    for j, value in enumerate(row):
                                        table.cell(i, j).text = str(value) if value is not None else ""
                            doc.add_paragraph()
                            if self.progressShow:
                                self.progressShow(int(yy * 100 / len(workbook.sheetnames)))
                        workbook.close()
                    elif ext == ".csv":
                        with open(ww, mode="r", encoding="utf-8") as csv_file:
                            csv_reader = csv.reader(csv_file)
                            table_data = list(csv_reader)
                        if table_data:
                            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                            for i, row in enumerate(table_data):
                                for j, value in enumerate(row):
                                    table.cell(i, j).text = str(value) if value is not None else ""
                                if self.progressShow:
                                    self.progressShow(int(i * 100 / len(table_data)))
                    doc.save(_[xx])
                    __.append(mf.setReturn(True, [type(_[xx])], [_[xx]]))
                except Exception as e:
                    __.append(mf.setReturn(False, [type(e)], [e]))
            else:
                for i in __:
                    try:
                        os.remove(mf.Return(i).getValues(0))
                    except:pass
                break
        if self.progressShow:
            self.progressShow(100)
        return __

    def toWORD(self, path):
        temp = self.Convert(self.file, [])
        path = path if type(path) == list else path
        _ = list()
        if not self.stop:
            self.toKill = WordToWord([mf.Return(ww).getValues(0) for ww in temp])
            _ = self.toKill.toWORD(path)
        if self.progressShow:
            self.progressShow(100, True, self.stop, _)
        return _ if not self.stop else []

    def kill(self):
        self.stop = True
        if self.toKill:
            try:
                self.toKill.kill()
            except Exception as e:
                print(e)


class ImgToWord:
    def __init__(self, file):
        self.file = file if type(file) == list else [file]
        self.stop = False
        self.toKill = None
        self.progressShow = None
    
    def toWord(self, path):
        _ = list()
        for ww in self.file:
            if not self.stop:
                self.toKill = MixedToWord([["image",  ww, None]])
                self.toKill.progressShow = self.progressShow
                _.append(self.toKill.toWORD(path))
            else:
                for i in _:
                    try:
                        os.remove(mf.Return(i).getValues(0))
                    except:pass
                _=[]
                break
        if self.progressShow:
            self.progressShow(100, True, self.stop, _)
        return _
    
    def kill(self):
        self.stop = True
        if self.toKill:
            try:
                self.toKill.kill()
            except Exception as e:
                print(e)

